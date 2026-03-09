import json
import os
from typing import List

from core.rag.interfaces import DocumentParser
from core.rag.models import ParsedDocument


class TextLikeParser(DocumentParser):
    def __init__(self, max_chars: int = 120000):
        self.max_chars = max(2000, int(max_chars))

    @staticmethod
    def _read_text_file(path: str) -> str:
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(path, "rb") as f:
            return f.read().decode("utf-8", errors="ignore")

    def _parse_csv(self, path: str) -> str:
        try:
            import pandas as pd
        except Exception:
            return self._read_text_file(path)

        df = pd.read_csv(path)
        head = df.head(200)
        table = head.to_string(index=False)
        return (
            f"CSV file: {os.path.basename(path)}\n"
            f"Rows: {len(df)} | Columns: {len(df.columns)}\n"
            f"Column names: {', '.join(map(str, df.columns.tolist()))}\n\n"
            f"{table}"
        )

    def _parse_json(self, path: str) -> str:
        raw = self._read_text_file(path)
        try:
            data = json.loads(raw)
        except Exception:
            return raw
        return json.dumps(data, indent=2, ensure_ascii=True)

    def _parse_docx(self, path: str) -> str:
        try:
            from docx import Document
        except Exception:
            return self._read_text_file(path)

        doc = Document(path)
        lines: List[str] = []
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if txt:
                lines.append(txt)
        return "\n".join(lines)

    def parse(self, path: str) -> ParsedDocument:
        ext = os.path.splitext(path)[1].lower()

        if ext == ".csv":
            text = self._parse_csv(path)
            source_type = "csv"
        elif ext == ".json":
            text = self._parse_json(path)
            source_type = "json"
        elif ext == ".docx":
            text = self._parse_docx(path)
            source_type = "docx"
        else:
            text = self._read_text_file(path)
            source_type = ext.lstrip(".") or "text"

        text = (text or "").strip()
        if len(text) > self.max_chars:
            text = text[: self.max_chars]

        return ParsedDocument(
            source=path,
            source_type=source_type,
            pages=[text],
            metadata={"source_type": source_type, "extension": ext},
        )
