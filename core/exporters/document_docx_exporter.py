import os
from typing import Dict, List


class DocumentDocxExporter:
    def export(
        self,
        *,
        title: str,
        outline: Dict[str, object],
        text: str,
        output_path: str,
        chunks: List[Dict[str, object]] | None = None,
        charts: List[Dict[str, object]] | None = None,
    ) -> None:
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX export") from exc

        doc = Document()
        doc.add_heading(title or "Generated Document", level=0)

        sections = self._split_sections(text)
        if sections:
            for section_title, body in sections:
                doc.add_heading(section_title, level=1)
                for para in body.split("\n\n"):
                    para = para.strip()
                    if para:
                        doc.add_paragraph(para)
        else:
            for para in (text or "").split("\n\n"):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)

        chart_items = [c for c in (charts or []) if isinstance(c, dict) and c.get("image_path")]
        if chart_items:
            doc.add_page_break()
            doc.add_heading("Charts & Visuals", level=1)
            for idx, chart in enumerate(chart_items, start=1):
                image_path = str(chart.get("image_path", "")).strip()
                if not image_path or not os.path.exists(image_path):
                    continue
                title_text = str(chart.get("title", f"Chart {idx}")).strip() or f"Chart {idx}"
                caption = str(chart.get("caption", "")).strip()
                sources = chart.get("evidence_sources") or []
                source_text = ", ".join(
                    os.path.basename(str(s).strip()) or str(s).strip()
                    for s in sources
                    if str(s).strip()
                )

                doc.add_heading(f"{idx}. {title_text}", level=2)
                if caption:
                    doc.add_paragraph(caption)
                if source_text:
                    doc.add_paragraph(f"Sources: {source_text}")
                doc.add_picture(image_path, width=Inches(6.3))

        refs = self._format_references(chunks or [])
        if refs:
            doc.add_page_break()
            doc.add_heading("References", level=1)
            for line in refs:
                if not line:
                    doc.add_paragraph("")
                    continue
                if line.endswith(":"):
                    doc.add_heading(line[:-1], level=2)
                    continue
                doc.add_paragraph(line)

        doc.save(output_path)

    @staticmethod
    def _split_sections(text: str) -> List[tuple[str, str]]:
        lines = (text or "").splitlines()
        sections: List[tuple[str, str]] = []
        current_title = None
        current_body: List[str] = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                current_body.append("")
                continue

            if current_title is None:
                current_title = stripped
                continue

            if stripped and stripped == stripped.title() and len(stripped.split()) <= 8 and current_body:
                sections.append((current_title, "\n".join(current_body).strip()))
                current_title = stripped
                current_body = []
                continue

            current_body.append(line)

        if current_title:
            sections.append((current_title, "\n".join(current_body).strip()))
        return [(t, b) for t, b in sections if b]

    @staticmethod
    def _format_references(chunks: List[Dict[str, object]]) -> List[str]:
        grouped: Dict[str, List[str]] = {}
        for chunk in chunks:
            section = str(chunk.get("section_title", "Section"))
            citations = chunk.get("citations") or []
            if not isinstance(citations, list):
                continue
            for c in citations:
                if not isinstance(c, dict):
                    continue
                source = str(c.get("source", "unknown"))
                page = str(c.get("page", "?"))
                score = float(c.get("score", 0.0))
                grouped.setdefault(section, [])
                ref = f"- {source} (page {page}, score {score:.3f})"
                if ref not in grouped[section]:
                    grouped[section].append(ref)

        lines: List[str] = []
        for section, refs in grouped.items():
            lines.append(f"{section}:")
            lines.extend(refs[:10])
            lines.append("")
        return lines
